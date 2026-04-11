#!/usr/bin/env python3
"""Generate BakeShop POS User Manual as .docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ─── Styles ───
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x8B, 0x5E, 0x3C)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

def add_feature(title, desc):
    p = doc.add_paragraph()
    run = p.add_run(f"  {title}")
    run.bold = True
    run.font.size = Pt(11)
    p2 = doc.add_paragraph(desc)
    p2.paragraph_format.left_indent = Cm(1)
    p2.paragraph_format.space_after = Pt(8)

# ═══════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("BakeShop POS")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x8B, 0x5E, 0x3C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Panduan Pengguna / User Manual")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xA0, 0x67, 0x3C)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("Aplikasi Point of Sale untuk Toko Bahan Kue")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run(f"Versi 1.0 — {datetime.date.today().strftime('%d %B %Y')}")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════
doc.add_heading("Daftar Isi", level=1)
toc_items = [
    "1. Pendahuluan",
    "2. Akun Demo & Login",
    "3. Struktur Role & Hak Akses",
    "4. Halaman Login",
    "5. Beranda (Dashboard)",
    "6. Kasir (POS)",
    "7. Stok (Inventory)",
    "8. Pesanan (Orders)",
    "9. Pengaturan (Settings)",
    "10. Fitur Umum",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 1. PENDAHULUAN
# ═══════════════════════════════════════════════════════
doc.add_heading("1. Pendahuluan", level=1)
doc.add_paragraph(
    "BakeShop POS adalah aplikasi Point of Sale (POS) berbasis web yang dirancang khusus untuk toko bahan kue. "
    "Aplikasi ini mendukung pengelolaan transaksi penjualan, manajemen stok & batch (FIFO), pencatatan pesanan, "
    "manajemen tim, dan pengaturan toko — semuanya dalam satu platform yang mudah digunakan."
)
doc.add_paragraph("Fitur utama meliputi:")
features_intro = [
    "Penjualan kasir dengan barcode scanner, diskon item/pesanan, dan multi-metode pembayaran",
    "Manajemen stok dengan tracking batch FIFO dan peringatan kedaluwarsa",
    "Sistem role-based access control (RBAC) dengan 5 level peran",
    "Kas awal & tutup kasir dengan penghitungan selisih",
    "Bilingual (Bahasa Indonesia & English) dengan mode gelap/terang",
    "Cetak struk dan label barcode",
    "Ekspor data ke CSV dan Excel",
]
for f in features_intro:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 2. AKUN DEMO & LOGIN
# ═══════════════════════════════════════════════════════
doc.add_heading("2. Akun Demo & Login", level=1)
doc.add_paragraph(
    "Aplikasi menyediakan akun demo untuk pengujian. Semua akun demo menggunakan password yang sama. "
    "Staf baru yang didaftarkan melalui menu Pengaturan > Tim akan mendapatkan password default: bakeshop123"
)

doc.add_heading("Daftar Akun Demo", level=2)
add_table(
    ["Nama", "Email", "Password", "Role", "Deskripsi"],
    [
        ["Rina Wijaya", "rina@bakeshop.id", "admin", "Super Admin", "Akses penuh ke semua fitur"],
        ["Andi Pratama", "andi@bakeshop.id", "admin", "Admin / Pemilik", "Akses penuh ke semua fitur"],
        ["Siti Nurhaliza", "siti@bakeshop.id", "admin", "Kasir", "Akses kasir (POS) & pengaturan"],
        ["Dewi Lestari", "dewi@bakeshop.id", "admin", "Staff", "Akses beranda, stok, & pengaturan"],
        ["Budi Santoso", "budi@bakeshop.id", "admin", "User", "Akses beranda, pesanan, & pengaturan"],
    ]
)
doc.add_paragraph()
doc.add_paragraph(
    "Catatan: Di halaman login tersedia tombol Quick Login untuk akun Admin, Kasir, dan Staff "
    "sehingga tidak perlu mengetik email dan password secara manual saat demo."
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 3. ROLE & HAK AKSES
# ═══════════════════════════════════════════════════════
doc.add_heading("3. Struktur Role & Hak Akses", level=1)
doc.add_paragraph(
    "Setiap pengguna memiliki role yang menentukan halaman mana saja yang dapat diakses. "
    "Berikut adalah matriks hak akses per role:"
)

add_table(
    ["Role", "Beranda", "Kasir (POS)", "Stok", "Pesanan", "Pengaturan"],
    [
        ["Super Admin", "Ya", "Ya", "Ya", "Ya", "Ya (penuh)"],
        ["Admin / Pemilik", "Ya", "Ya", "Ya", "Ya", "Ya (penuh)"],
        ["Kasir", "-", "Ya", "-", "-", "Ya (preferensi)"],
        ["Staff", "Ya", "-", "Ya", "-", "Ya (preferensi)"],
        ["User", "Ya", "-", "-", "Ya", "Ya (preferensi)"],
    ]
)
doc.add_paragraph()
doc.add_paragraph("Keterangan tambahan:")
notes = [
    "Super Admin & Admin: Dapat mengelola toko, tim, rekening bank, pajak, dan melihat log aktivitas di Pengaturan.",
    "Kasir: Hanya melihat halaman POS dan Pengaturan (preferensi). Wajib membuka kasir (Buka Kasir) sebelum transaksi. Bisa melihat riwayat pesanan melalui tombol di halaman POS.",
    "Staff: Fokus pada operasional stok — dapat menambah/mengeluarkan stok, mengelola pemasok, dan melihat peringatan kedaluwarsa. Tidak dapat mengakses POS atau pesanan.",
    "User: Hanya bisa melihat pesanan dan pengaturan preferensi. Tidak dapat melakukan transaksi atau mengelola stok.",
    "Stok Write Access: Hanya Super Admin, Admin, dan Staff yang dapat menambah/mengeluarkan stok (role lain hanya bisa melihat)."
]
for n in notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 4. HALAMAN LOGIN
# ═══════════════════════════════════════════════════════
doc.add_heading("4. Halaman Login", level=1)
doc.add_paragraph(
    "Halaman login adalah halaman pertama yang muncul saat membuka aplikasi. "
    "Pengguna harus memasukkan email dan password untuk masuk ke sistem."
)

doc.add_heading("Fitur di Halaman Login", level=2)

add_feature("Login dengan Email & Password",
    "Masukkan alamat email dan password yang terdaftar, lalu klik tombol 'Masuk'. "
    "Jika kredensial salah, akan muncul pesan error 'Email atau sandi salah'. "
    "Jika akun dinonaktifkan oleh admin, akan muncul pesan 'Akun nonaktif'.")

add_feature("Quick Login (Demo)",
    "Di bagian bawah form login terdapat tombol-tombol quick login untuk akun demo "
    "(Admin, Kasir, Staff). Cukup klik salah satu tombol untuk langsung masuk tanpa mengetik email/password.")

add_feature("Ganti Tema & Bahasa",
    "Di atas form login terdapat tombol untuk mengganti tema (terang/gelap) dan bahasa (English/Indonesia). "
    "Pengaturan ini akan tersimpan dan berlaku setelah login.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 5. BERANDA (DASHBOARD)
# ═══════════════════════════════════════════════════════
doc.add_heading("5. Beranda (Dashboard)", level=1)
doc.add_paragraph(
    "Halaman Beranda menampilkan ringkasan informasi harian yang disesuaikan berdasarkan role pengguna. "
    "Halaman ini tersedia untuk: Super Admin, Admin, Staff, dan User."
)

doc.add_heading("5.1 Tampilan Owner / Admin", level=2)
add_feature("Kartu Statistik",
    "Menampilkan total pendapatan (revenue), jumlah pesanan, dan jumlah peringatan stok rendah dalam bentuk kartu visual yang mudah dibaca.")
add_feature("Transaksi Terakhir",
    "Daftar 4 pesanan terbaru dengan ID pesanan, nama pelanggan, waktu, dan total. "
    "Klik pesanan untuk melihat detail lengkap termasuk item, diskon, dan bukti pembayaran.")
add_feature("Stok Menipis",
    "Menampilkan hingga 5 produk yang stoknya di bawah batas minimum. "
    "Ditandai dengan badge berwarna untuk stok rendah (oranye) dan stok habis (merah). "
    "Klik produk untuk melihat detail lengkap.")

doc.add_heading("5.2 Tampilan Staff", level=2)
add_feature("Kartu Statistik Staff",
    "Menampilkan jumlah total produk, peringatan stok rendah, dan jumlah batch mendekati kedaluwarsa.")
add_feature("Stok Menipis",
    "Sama seperti tampilan admin — daftar produk dengan stok di bawah minimum.")
add_feature("Peringatan Kedaluwarsa",
    "Menampilkan batch-batch produk yang akan kedaluwarsa dalam 60 hari ke depan. "
    "Ditandai dengan warna: merah (sudah kedaluwarsa), oranye (mendesak, < 14 hari), kuning (dalam 60 hari).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 6. KASIR (POS)
# ═══════════════════════════════════════════════════════
doc.add_heading("6. Kasir (POS)", level=1)
doc.add_paragraph(
    "Halaman Kasir (POS) adalah pusat transaksi penjualan. Tersedia untuk Super Admin, Admin, dan Kasir. "
    "Sebelum dapat memproses transaksi, pengguna WAJIB membuka kasir terlebih dahulu."
)

doc.add_heading("6.1 Buka Kasir (Open Register)", level=2)
add_feature("Wajib Sebelum Transaksi",
    "Saat pertama kali masuk ke halaman POS, sistem akan memblokir akses dan menampilkan layar "
    "'Kasir belum dibuka'. Pengguna harus mengklik 'Buka Kasir' dan memasukkan jumlah kas awal (modal uang tunai).")
add_feature("Input Kas Awal",
    "Masukkan jumlah uang tunai yang ada di laci kasir sebagai modal awal. "
    "Angka ini akan digunakan sebagai dasar perhitungan saat tutup kasir. "
    "Setelah dibuka, kas awal tidak bisa diedit — harus tutup kasir lalu buka sesi baru.")
add_feature("Riwayat Kasir Hari Ini",
    "Di bawah tombol Buka Kasir, ditampilkan riwayat sesi kasir yang sudah ditutup hari ini "
    "beserta informasi kas awal, kas diharapkan, kas aktual, dan selisih.")

doc.add_heading("6.2 Katalog Produk", level=2)
add_feature("Pencarian Produk",
    "Ketik nama produk atau SKU di kolom pencarian. Pencarian bersifat real-time (debounced). "
    "Tekan tombol '/' di keyboard untuk langsung fokus ke kolom pencarian.")
add_feature("Filter Kategori",
    "Gunakan pill kategori di bawah kolom pencarian untuk memfilter produk berdasarkan kategori "
    "(Tepung, Gula, Susu & Telur, Cokelat, dll.). Klik 'Semua' untuk menampilkan seluruh produk.")
add_feature("Kartu Produk",
    "Setiap produk ditampilkan dalam kartu dengan gambar, nama, harga, dan indikator stok. "
    "Badge stok berwarna: hijau (normal), oranye (rendah), merah (habis). "
    "Klik gambar/nama untuk melihat detail produk. Tombol '+ Satuan' dan 'Dus' untuk menambahkan ke keranjang.")
add_feature("Barcode Scanner",
    "Aplikasi mendukung barcode scanner hardware. Saat scanner memindai kode, produk dengan SKU yang sesuai "
    "otomatis ditambahkan ke keranjang. Jika SKU tidak ditemukan, muncul pesan error.")

doc.add_heading("6.3 Keranjang Belanja", level=2)
add_feature("Tampilan Desktop vs Mobile",
    "Desktop: Keranjang tampil sebagai panel tetap di sisi kanan. "
    "Mobile: Keranjang muncul sebagai floating bar di bawah, klik untuk membuka modal keranjang penuh.")
add_feature("Pelanggan / Member",
    "Kolom input untuk mencari member berdasarkan nomor telepon. "
    "Pilih member yang ada atau tambah member baru langsung dari sini (nama + nomor telepon).")
add_feature("Kelola Kuantitas",
    "Gunakan tombol +/- untuk mengubah jumlah setiap item. "
    "Sistem akan validasi stok — jika stok tidak cukup, muncul pesan 'Stok tidak cukup'. "
    "Klik ikon tempat sampah untuk menghapus item.")
add_feature("Diskon Item",
    "Klik ikon tag pada item untuk menambahkan diskon per item. "
    "Pilih tipe diskon: persen (%) atau nominal tetap (Rp). "
    "Diskon terlihat langsung pada subtotal item dengan harga coret.")
add_feature("Diskon Pesanan",
    "Klik 'Tambah Diskon' di bawah subtotal untuk memberikan diskon keseluruhan pesanan. "
    "Bisa berupa persen atau nominal tetap. Diskon pesanan diterapkan setelah diskon item.")
add_feature("Metode Pembayaran",
    "Pilih salah satu dari 4 metode: Tunai, Kartu, Transfer, atau QRIS. "
    "Metode yang aktif ditandai dengan warna gradien oranye.")
add_feature("Hapus Keranjang",
    "Tombol 'Hapus' mengosongkan seluruh keranjang termasuk pelanggan dan diskon.")

doc.add_heading("6.4 Checkout (Pembayaran)", level=2)
add_feature("Modal Checkout",
    "Klik 'Bayar' untuk membuka modal checkout. Menampilkan total pesanan "
    "lengkap dengan rincian subtotal, diskon, dan PPN.")
add_feature("Pembayaran Tunai",
    "Masukkan jumlah uang yang diterima. Sistem otomatis menghitung kembalian. "
    "Tombol 'Konfirmasi' hanya aktif jika uang diterima >= total pesanan.")
add_feature("Pembayaran Transfer",
    "Pilih rekening bank tujuan (jika sudah dikonfigurasi di Pengaturan). "
    "Upload bukti transfer/pembayaran wajib sebelum konfirmasi.")
add_feature("Pembayaran QRIS",
    "Upload bukti pembayaran QRIS wajib sebelum konfirmasi.")
add_feature("Struk Pesanan",
    "Setelah checkout berhasil, muncul modal struk dengan detail pesanan dan barcode. "
    "Klik 'Cetak Struk' untuk mencetak struk melalui printer.")

doc.add_heading("6.5 Tutup Kasir (Close Register)", level=2)
add_feature("Hitung Kas Akhir",
    "Klik tombol '$' (di sebelah pencarian) untuk membuka modal Tutup Kasir. "
    "Sistem menampilkan kas diharapkan = kas awal + total penjualan tunai selama sesi.")
add_feature("Input Kas Aktual",
    "Masukkan jumlah uang tunai yang sebenarnya ada di laci kasir. "
    "Sistem otomatis menghitung selisih dan menampilkan status: "
    "Seimbang (hijau), Lebih (biru), atau Kurang (merah).")
add_feature("Catatan",
    "Kolom opsional untuk mencatat alasan jika ada selisih kas.")
add_feature("Riwayat Kasir",
    "Di bawah form, ditampilkan riwayat sesi kasir yang sudah ditutup hari ini.")

doc.add_heading("6.6 Riwayat Pesanan (Order History)", level=2)
add_feature("Tombol Riwayat",
    "Klik ikon dokumen di sebelah pencarian untuk membuka modal Riwayat Pesanan. "
    "Menampilkan semua pesanan yang dibuat selama sesi kasir aktif saat ini.")
add_feature("Detail Pesanan",
    "Setiap pesanan menampilkan ID, status (badge berwarna), waktu, pelanggan, total, "
    "dan daftar item yang dibeli.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 7. STOK (INVENTORY)
# ═══════════════════════════════════════════════════════
doc.add_heading("7. Stok (Inventory)", level=1)
doc.add_paragraph(
    "Halaman Stok adalah pusat manajemen inventaris. Tersedia untuk Super Admin, Admin, dan Staff. "
    "Halaman ini memiliki 6 tab: Ringkasan, Masuk, Keluar, Kedaluwarsa, Riwayat, dan Pemasok."
)

doc.add_heading("7.1 Tab Ringkasan (Overview)", level=2)
add_feature("Statistik Cepat",
    "Menampilkan 3 angka penting: jumlah stok rendah, stok habis, dan produk nonaktif.")
add_feature("Filter Produk",
    "Filter berdasarkan status: Semua, Stok Rendah, Stok Habis, atau Nonaktif. "
    "Pencarian produk berdasarkan nama atau SKU.")
add_feature("Daftar Produk",
    "Setiap produk menampilkan gambar, nama, SKU, stok saat ini, dan status (Normal/Rendah/Habis). "
    "Klik produk untuk melihat detail lengkap (harga beli/jual, margin, batch, pergerakan terakhir).")
add_feature("Tampilkan/Sembunyikan di POS",
    "Toggle untuk menampilkan atau menyembunyikan produk dari halaman kasir. "
    "Produk yang disembunyikan ditandai badge 'Nonaktif'.")
add_feature("Cetak Label Barcode",
    "Klik ikon barcode pada setiap produk untuk mencetak label barcode (SKU + nama produk).")
add_feature("Tambah Produk Baru",
    "Tombol '+ Tambah Produk' membuka form untuk membuat produk baru: "
    "nama (EN/ID), SKU, kategori, harga beli, harga jual, isi per dus, satuan, stok minimum, dan gambar.")
add_feature("Edit Produk",
    "Tombol edit (pensil) pada setiap produk membuka form edit untuk mengubah "
    "nama, SKU, kategori, harga, isi per dus, satuan, gambar, dan stok minimum.")
add_feature("Tambah Kategori",
    "Tombol 'Tambah Kategori' membuka form untuk membuat kategori baru dengan nama (EN/ID) dan warna.")
add_feature("Ekspor Data Produk",
    "Tombol ekspor CSV dan Excel untuk mengunduh daftar produk beserta stok dan harga.")

doc.add_heading("7.2 Tab Stok Masuk (Stock In)", level=2)
add_feature("Catat Stok Masuk",
    "Tombol '+ Stok Masuk' membuka form pencatatan: pilih produk, jumlah, tipe unit (satuan/dus), "
    "harga satuan, catatan, dan tanggal kedaluwarsa (opsional).")
add_feature("Pemasok & Termin",
    "Pilih pemasok untuk stok masuk (opsional). Jika dipilih, atur termin pembayaran "
    "(COD, Net 30/60/90 hari) dan status pembayaran (Lunas/Belum Lunas).")
add_feature("Daftar Stok Masuk",
    "Riwayat semua stok masuk dikelompokkan per tanggal, menampilkan produk, jumlah, "
    "harga, pemasok, termin pembayaran, dan status bayar.")
add_feature("Tandai Lunas",
    "Tombol 'Tandai Lunas' untuk mengubah status pembayaran dari Belum Lunas ke Lunas.")

doc.add_heading("7.3 Tab Stok Keluar (Stock Out)", level=2)
add_feature("Catat Stok Keluar",
    "Form pencatatan stok keluar untuk barang rusak, retur, atau penyesuaian. "
    "Pilih produk, jumlah, tipe unit, harga, dan alasan.")
add_feature("Daftar Stok Keluar",
    "Riwayat semua stok keluar dikelompokkan per tanggal.")

doc.add_heading("7.4 Tab Kedaluwarsa (Expiry)", level=2)
add_feature("Peringatan Kedaluwarsa",
    "Menampilkan semua batch produk yang akan kedaluwarsa dalam 90 hari ke depan. "
    "Ditandai dengan warna berdasarkan urgensi: merah (sudah kedaluwarsa), "
    "oranye (< 14 hari), kuning (< 90 hari). "
    "Menampilkan nomor batch, jumlah sisa, tanggal kedaluwarsa, dan countdown hari.")

doc.add_heading("7.5 Tab Riwayat (History)", level=2)
add_feature("Riwayat Pergerakan Stok",
    "Gabungan semua stok masuk dan keluar dikelompokkan per tanggal. "
    "Setiap entri menampilkan arah (masuk/keluar), produk, jumlah, harga, dan catatan.")

doc.add_heading("7.6 Tab Pemasok (Suppliers)", level=2)
add_feature("Daftar Pemasok",
    "Menampilkan semua pemasok yang terdaftar dengan nama, nomor telepon, dan email. "
    "Klik pemasok untuk melihat detail lengkap termasuk alamat dan riwayat stok masuk.")
add_feature("Tambah Pemasok",
    "Form untuk menambah pemasok baru: nama, telepon, email, dan alamat.")
add_feature("Edit Pemasok",
    "Ubah informasi pemasok yang sudah ada.")
add_feature("Hapus Pemasok",
    "Hapus pemasok dengan konfirmasi. Penghapusan bersifat permanen.")
add_feature("Tagihan Belum Lunas",
    "Di bawah daftar pemasok, ditampilkan daftar tagihan yang belum lunas dan sudah lewat jatuh tempo.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 8. PESANAN (ORDERS)
# ═══════════════════════════════════════════════════════
doc.add_heading("8. Pesanan (Orders)", level=1)
doc.add_paragraph(
    "Halaman Pesanan menampilkan seluruh riwayat transaksi. "
    "Tersedia untuk Super Admin, Admin, dan User."
)

doc.add_heading("8.1 Filter & Pencarian", level=2)
add_feature("Filter Periode",
    "Pill filter periode: Hari Ini, Kemarin, Minggu Ini, Bulan Ini, atau Semua. "
    "Statistik ringkasan (pendapatan, jumlah pesanan, rata-rata, batal) otomatis mengikuti periode yang dipilih.")
add_feature("Pencarian",
    "Cari pesanan berdasarkan ID pesanan atau nama pelanggan.")
add_feature("Filter Status",
    "Filter berdasarkan status: Semua, Selesai, Tertunda, Dibatalkan, atau Direfund.")

doc.add_heading("8.2 Daftar Pesanan", level=2)
add_feature("Kartu Pesanan",
    "Setiap pesanan menampilkan ID, pelanggan, waktu, total, status, metode pembayaran, "
    "dan daftar item yang dibeli. Klik pesanan untuk melihat detail lengkap.")
add_feature("Detail Pesanan (Modal)",
    "Modal detail menampilkan: daftar item beserta diskon per item, "
    "subtotal, diskon pesanan, PPN, total, metode pembayaran, kasir yang memproses, "
    "dan bukti pembayaran (jika transfer/QRIS). "
    "Tombol 'Batalkan Pesanan' untuk membatalkan (stok dikembalikan). "
    "Tombol 'Refund' untuk memproses refund penuh atau sebagian.")
add_feature("Pagination",
    "Pesanan ditampilkan 20 per halaman. Klik 'Tampilkan lagi' untuk memuat lebih banyak.")

doc.add_heading("8.3 Ekspor & Cetak", level=2)
add_feature("Ekspor CSV",
    "Unduh data pesanan dalam format CSV.")
add_feature("Ekspor Excel",
    "Unduh data pesanan dalam format Excel (.xlsx).")
add_feature("Cetak Laporan",
    "Cetak laporan transaksi untuk periode yang dipilih melalui printer.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 9. PENGATURAN (SETTINGS)
# ═══════════════════════════════════════════════════════
doc.add_heading("9. Pengaturan (Settings)", level=1)
doc.add_paragraph(
    "Halaman Pengaturan tersedia untuk semua role, namun fitur yang tampil berbeda sesuai hak akses. "
    "Admin/Super Admin melihat 4 tab, sedangkan role lain hanya melihat tab Preferensi."
)

doc.add_heading("9.1 Tab Preferensi (Semua Role)", level=2)
add_feature("Tema Tampilan",
    "Pilih antara mode Terang atau Gelap. Pengaturan tersimpan secara permanen di browser.")
add_feature("Bahasa",
    "Pilih antara English atau Indonesia. Seluruh teks aplikasi berubah sesuai bahasa yang dipilih.")
add_feature("Keluar (Sign Out)",
    "Tombol untuk keluar dari akun saat ini dan kembali ke halaman login.")

doc.add_heading("9.2 Tab Toko (Admin Only)", level=2)
add_feature("Informasi Toko",
    "Atur nama toko, alamat, dan nomor telepon. Informasi ini digunakan pada struk pembayaran.")
add_feature("Pajak (PPN)",
    "Atur tarif PPN (0-100%). PPN otomatis ditambahkan ke setiap transaksi penjualan. "
    "Set 0% untuk menonaktifkan PPN.")
add_feature("Rekening Bank",
    "Kelola daftar rekening bank yang tampil saat pembayaran transfer di kasir. "
    "Tambah rekening dengan memilih bank dari daftar 50+ bank Indonesia, "
    "lalu isi nomor rekening dan nama pemilik. Rekening bisa dihapus dengan konfirmasi.")

doc.add_heading("9.3 Tab Tim (Admin Only)", level=2)
add_feature("Daftar Staf",
    "Menampilkan semua pengguna dengan role Kasir dan Staff dalam bentuk kartu. "
    "Setiap kartu menunjukkan nama, email, telepon, dan role. "
    "Akun yang dinonaktifkan ditandai badge merah 'Akun nonaktif'.")
add_feature("Lihat Detail Staf",
    "Klik kartu staf untuk expand dan melihat detail lengkap: NIK, nomor telepon, "
    "tanggal lahir, dan email.")
add_feature("Daftarkan Staf Baru",
    "Tombol 'Daftar Staf' membuka form pendaftaran: nama lengkap, NIK, email, "
    "nomor telepon, tanggal lahir, dan pilihan role (Kasir atau Staff). "
    "Password default otomatis: bakeshop123")
add_feature("Reset Password",
    "Di detail staf, tersedia form untuk mereset password. "
    "Masukkan password baru dan klik Simpan. Perubahan tercatat di log aktivitas.")
add_feature("Aktifkan / Nonaktifkan Akun",
    "Tombol toggle untuk mengaktifkan atau menonaktifkan akun staf. "
    "Akun nonaktif tidak bisa login ke aplikasi. "
    "Tampilan kartu staf menjadi semi-transparan saat nonaktif.")
add_feature("Hapus Staf",
    "Tombol hapus dengan konfirmasi dua langkah (klik Hapus, lalu Konfirmasi). "
    "Penghapusan bersifat permanen dan tercatat di log aktivitas.")

doc.add_heading("9.4 Tab Aktivitas (Admin Only)", level=2)
add_feature("Log Aktivitas",
    "Riwayat kronologis semua aktivitas penting di aplikasi: "
    "pesanan dibuat, pesanan dibatalkan, pesanan direfund, stok disesuaikan, "
    "produk ditambah/diedit, pengaturan diubah, pengguna didaftarkan, "
    "akun diaktifkan/dinonaktifkan, akun dihapus, password direset, "
    "kasir dibuka/ditutup. "
    "Setiap entri menampilkan jenis aktivitas (berwarna), detail, waktu, dan nama pengguna.")
add_feature("Pagination",
    "Entri ditampilkan 20 per halaman. Klik 'Tampilkan lagi' untuk memuat lebih banyak.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 10. FITUR UMUM
# ═══════════════════════════════════════════════════════
doc.add_heading("10. Fitur Umum", level=1)

doc.add_heading("10.1 Navigasi", level=2)
add_feature("Bottom Navigation",
    "Bar navigasi di bagian bawah layar menampilkan menu yang tersedia sesuai role pengguna. "
    "Menu aktif ditandai dengan ikon dan teks berwarna aksen. "
    "Navigasi keyboard: gunakan tombol panah kiri/kanan untuk berpindah tab.")
add_feature("Header",
    "Bar header di bagian atas menampilkan logo BakeShop, nama halaman aktif, "
    "nama pengguna, dan role yang sedang login.")

doc.add_heading("10.2 Detail Produk", level=2)
add_feature("Modal Detail Produk",
    "Klik produk di manapun (beranda, kasir, stok, pesanan) untuk membuka modal detail. "
    "Menampilkan: gambar produk, nama (EN/ID), SKU, kategori, harga beli & jual, "
    "margin keuntungan, stok saat ini, informasi batch (FIFO), dan pergerakan stok terakhir.")

doc.add_heading("10.3 Detail Pesanan", level=2)
add_feature("Modal Detail Pesanan",
    "Klik pesanan untuk membuka modal detail. "
    "Menampilkan: ID, status, tanggal, pelanggan, kasir, daftar item, diskon, PPN, total, "
    "metode pembayaran, dan bukti pembayaran. "
    "Tombol aksi: Batalkan Pesanan (stok dikembalikan) dan Refund (penuh atau sebagian).")

doc.add_heading("10.4 Cetak", level=2)
add_feature("Cetak Struk",
    "Setelah checkout berhasil, klik 'Cetak Struk' pada modal struk untuk mencetak. "
    "Struk berisi: info toko, daftar item, diskon, PPN, total, metode pembayaran, dan barcode pesanan.")
add_feature("Cetak Label Barcode",
    "Di halaman Stok (tab Ringkasan), klik ikon barcode pada produk untuk mencetak label "
    "berisi barcode SKU dan nama produk.")
add_feature("Cetak Laporan",
    "Di halaman Pesanan, klik 'Cetak Laporan' untuk mencetak ringkasan transaksi "
    "berdasarkan periode yang dipilih.")

doc.add_heading("10.5 Notifikasi", level=2)
add_feature("Toast Notification",
    "Notifikasi singkat muncul di bagian atas tengah layar untuk mengkonfirmasi aksi: "
    "pesanan berhasil, stok tercatat, produk ditambahkan, pengaturan disimpan, dll. "
    "Notifikasi otomatis hilang setelah 2.5 detik.")

doc.add_heading("10.6 Responsif", level=2)
add_feature("Mobile & Desktop",
    "Aplikasi dirancang mobile-first dan responsive. "
    "Di mobile: keranjang tampil sebagai floating bar, modal slide up dari bawah. "
    "Di desktop: keranjang tampil sebagai panel tetap di sisi kanan, layout lebih luas.")

# ─── Save ───
output_path = "/Users/liaseptiany/Documents/web-freelance/bakeshop-fe/BakeShop_POS_User_Manual.docx"
doc.save(output_path)
print(f"Manual saved to: {output_path}")
